#!/usr/bin/env python3
"""文档解析脚本 - 支持 PDF、DOCX、TXT"""

import json
import sys
import os
from pathlib import Path


def parse_pdf(filepath):
    """解析 PDF 文件"""
    from PyPDF2 import PdfReader
    reader = PdfReader(filepath)
    
    text_parts = []
    structure = []
    tables = []
    
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        text_parts.append(f"--- 第 {i+1} 页 ---\n{page_text}")
        structure.append({
            "type": "page",
            "number": i + 1,
            "text": page_text[:200] + "..." if len(page_text) > 200 else page_text,
        })
    
    return {
        "text": "\n\n".join(text_parts),
        "tables": tables,
        "structure": structure,
        "pages": len(reader.pages),
        "format": "pdf",
    }


def parse_docx(filepath):
    """解析 DOCX 文件"""
    import docx
    doc = docx.Document(filepath)
    
    text_parts = []
    structure = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            text_parts.append(para.text)
            structure.append({
                "type": "paragraph",
                "style": style,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
            })
    
    # 表格
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "text": "\n\n".join(text_parts),
        "tables": tables,
        "structure": structure,
        "paragraphs": len(text_parts),
        "format": "docx",
    }


def parse_txt(filepath):
    """解析 TXT 文件"""
    text = Path(filepath).read_text(encoding="utf-8")
    
    lines = text.split("\n")
    structure = []
    for i, line in enumerate(lines):
        if line.strip():
            structure.append({
                "type": "line",
                "number": i + 1,
                "text": line[:100] + ("..." if len(line) > 100 else ""),
            })
    
    return {
        "text": text,
        "tables": [],
        "structure": structure,
        "lines": len(lines),
        "format": "txt",
    }


def parse_document(filepath):
    """统一解析入口"""
    filepath = str(filepath)
    ext = os.path.splitext(filepath)[1].lower()
    
    if not os.path.exists(filepath):
        return {"error": f"文件不存在: {filepath}"}
    
    try:
        if ext == ".pdf":
            return parse_pdf(filepath)
        elif ext == ".docx":
            return parse_docx(filepath)
        elif ext in (".txt", ".md"):
            return parse_txt(filepath)
        else:
            return {"error": f"不支持的格式: {ext}"}
    except Exception as e:
        return {"error": f"解析失败: {e}"}


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: parse_document.py <filepath>"}))
        sys.exit(1)
    
    result = parse_document(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
